from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.shared import Inches
# from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches, Pt, Cm
import os
import time


def generate_doc(page_num, outcome, file_path):
    """

    :return:
    """
    # 转换数据
    result = convert_data(page_num, outcome)

    # 打开文档
    document = Document()

    # 设置字体
    document.styles['Normal'].font.name = u'宋体'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    # 头部
    paragraph = document.add_paragraph()
    # 设置字号
    run = paragraph.add_run(u'技术需求匹配结果')
    run.font.size = Pt(24)
    # paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 段落文字居中设置
    # 设置粗体
    run = paragraph.add_run(u'').bold = True

    # 添加标题1： 技术需求描述
    paragraph = document.add_paragraph()
    # 设置字号
    run = paragraph.add_run(u'1. 技术需求来源')
    run.font.size = Pt(16)

    # 技术需求文本
    paragraph = document.add_paragraph()
    # 设置字号
    run = paragraph.add_run(u'请填写公司名')
    p_format = paragraph.paragraph_format
    p_format.first_line_indent = Inches(0.2)  # 首行缩进
    run.font.size = Pt(12)

    # 添加标题2： 技术需求描述
    paragraph = document.add_paragraph()
    # 设置字号
    run = paragraph.add_run(u'2. 技术需求描述')
    run.font.size = Pt(16)

    # 技术需求文本
    paragraph = document.add_paragraph()
    # 设置字号
    run = paragraph.add_run(u' 请填写技术需求描述')  # TODO： 缩进？？
    p_format = paragraph.paragraph_format
    p_format.first_line_indent = Inches(0.2)  # 首行缩进
    run.font.size = Pt(12)

    # 添加标题2： 技术需求匹配结果
    paragraph = document.add_paragraph()
    # 设置字号
    run = paragraph.add_run(u'3. 技术需求匹配结果')
    run.font.size = Pt(16)

    # 生成表格
    table = document.add_table(rows=4, cols=2, style='Table Grid')

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '专家团队'
    hdr_cells[0].paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells[1].text = result["member_str"]  # TODO： 替換

    hdr_cells = table.rows[1].cells
    hdr_cells[0].text = "所属院校"
    hdr_cells[0].paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells[1].text = result["school_institution_str"]

    hdr_cells = table.rows[2].cells
    hdr_cells[0].text = '相关专利'
    hdr_cells[0].paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells[1].text = result["patent_str"]

    hdr_cells = table.rows[3].cells
    hdr_cells[0].text = '相关项目'
    hdr_cells[0].paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells[1].text = result["project_str"]

    # 调整列宽
    table.cell(0, 0).width = Inches(2)
    table.cell(0, 1).width = Inches(8)

    # 保存文件
    cur_time = time.time()
    file_name = "result_" + str(cur_time) + ".docx"
    document.save(os.path.join(file_path, file_name))

    return file_name


def convert_data(page_num, outcome_dict):
    """
    转换 生成word 所需的数据
    :param page_num:
    :param outcome_dict:
    :return:
    """
    teacher_basic_info = outcome_dict["teacher_basic_info"]
    team = outcome_dict["team_list"][int(page_num)]
    school = team["school"]
    institution = team["institution"]
    lab = team["lab"]
    member_id_list = team["member_id_list"]
    patent_id_list = team["patent_id_list"]
    project_list = team["project_list"]
    patent_info = outcome_dict["patent_info"]
    member_str = ""
    for member_id in member_id_list:
        member_str += teacher_basic_info[member_id]["name"] + ","
    member_str = member_str[0:-1]
    school_institution_str = school + "-" + institution
    patent_str = ""
    i = 0
    for patent_id in patent_id_list:
        if i > 8:
            break
        i += 1
        patent_str += patent_info[patent_id][0] + "[" + patent_info[patent_id][1] + "]" + '\n'
    patent_str = patent_str[0:-1]
    project_str = ""
    i = 0
    for project in project_list:
        if i > 10:
            break
        i += 1
        project_str += project + '\n'
    project_str = project_str[0:-1]
    return {
        "member_str": member_str,
        "school_institution_str": school_institution_str,
        "patent_str": patent_str,
        "project_str": project_str,
    }








